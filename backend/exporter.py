import io
import re
from datetime import datetime
from reportlab.lib.pagesizes import letter

from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from docx import Document as DocxDocument
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def parse_markdown_to_reportlab(md_text: str, styles) -> list:
    """
    Very simple markdown parser to convert basic elements (*, **, #) into ReportLab Paragraphs.
    """
    flowables = []
    lines = md_text.split('\n')
    
    in_list = False
    
    # Custom styles
    normal_style = styles['Normal']
    body_style = ParagraphStyle(
        'CustomBody',
        parent=normal_style,
        fontSize=10,
        leading=14,
        textColor=colors.HexColor('#1f2937'),
        spaceAfter=6
    )
    h1_style = ParagraphStyle(
        'CustomH1',
        parent=styles['Heading1'],
        fontSize=16,
        leading=20,
        textColor=colors.HexColor('#1e3a8a'),
        spaceBefore=12,
        spaceAfter=8,
        keepWithNext=True
    )
    h2_style = ParagraphStyle(
        'CustomH2',
        parent=styles['Heading2'],
        fontSize=12,
        leading=16,
        textColor=colors.HexColor('#0f766e'),
        spaceBefore=10,
        spaceAfter=6,
        keepWithNext=True
    )
    quote_style = ParagraphStyle(
        'CustomQuote',
        parent=normal_style,
        fontSize=9,
        leading=13,
        textColor=colors.HexColor('#4b5563'),
        leftIndent=15,
        spaceAfter=8,
        backColor=colors.HexColor('#f3f4f6'),
        borderColor=colors.HexColor('#d1d5db'),
        borderWidth=1,
        borderPadding=5
    )
    
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            if in_list:
                in_list = False
            continue
            
        # Headers
        if line_strip.startswith('# '):
            text = line_strip[2:]
            text = clean_formatting(text)
            flowables.append(Paragraph(text, h1_style))
        elif line_strip.startswith('## '):
            text = line_strip[3:]
            text = clean_formatting(text)
            flowables.append(Paragraph(text, h2_style))
        elif line_strip.startswith('### '):
            text = line_strip[4:]
            text = clean_formatting(text)
            flowables.append(Paragraph(text, h2_style))
            
        # Blockquotes/Alerts
        elif line_strip.startswith('>'):
            text = line_strip.lstrip('> ').replace('[!NOTE]', '').replace('[!WARNING]', '').replace('[!IMPORTANT]', '')
            text = clean_formatting(text)
            flowables.append(Paragraph(text, quote_style))
            
        # Bullet list items
        elif line_strip.startswith('- ') or line_strip.startswith('* '):
            in_list = True
            text = line_strip[2:]
            text = clean_formatting(text)
            bullet_text = f"&bull; {text}"
            bullet_style = ParagraphStyle(
                'BulletStyle',
                parent=body_style,
                leftIndent=15,
                firstLineIndent=-10,
                spaceAfter=4
            )
            flowables.append(Paragraph(bullet_text, bullet_style))
            
        # Normal text paragraph
        else:
            text = clean_formatting(line_strip)
            # Check if it looks like bold key-value
            if text.startswith('<b>') and text.endswith('</b>') and len(text) < 100:
                # Keep together or bold heading
                flowables.append(Paragraph(text, ParagraphStyle('BoldLine', parent=body_style, fontName='Helvetica-Bold')))
            else:
                flowables.append(Paragraph(text, body_style))
                
    return flowables

def clean_formatting(text: str) -> str:
    # Convert markdown double asterisks to <b> and </b> tags
    text = re.sub(r'\*\*(.*?)\*\*', r'<b>\1</b>', text)
    # Convert single asterisks to <i> and </i> tags
    text = re.sub(r'\*(.*?)\*', r'<i>\1</i>', text)
    # Strip links [Text](URL)
    text = re.sub(r'\[(.*?)\]\(.*?\)', r'\1', text)
    return text

def generate_pdf_report(case_data: dict, analysis_output: str) -> bytes:
    buffer = io.BytesIO()
    
    # Document margins setup
    doc = SimpleDocTemplate(
        buffer,
        pagesize=letter,
        rightMargin=40,
        leftMargin=40,
        topMargin=40,
        bottomMargin=40
    )
    
    styles = getSampleStyleSheet()
    story = []
    
    # Official Header
    title_style = ParagraphStyle(
        'DocTitle',
        parent=styles['Heading1'],
        fontName='Helvetica-Bold',
        fontSize=20,
        leading=24,
        textColor=colors.HexColor('#0f172a'),
        alignment=1, # Center
        spaceAfter=4
    )
    subtitle_style = ParagraphStyle(
        'DocSubTitle',
        parent=styles['Normal'],
        fontSize=9,
        leading=12,
        textColor=colors.HexColor('#64748b'),
        alignment=1,
        spaceAfter=15
    )
    
    story.append(Paragraph("CRIMEGPT CASE REPORT & AI ANALYSIS", title_style))
    story.append(Paragraph("GENERATED VIA LAW ENFORCEMENT LEGAL INTELLIGENCE PORTAL", subtitle_style))
    story.append(Spacer(1, 5))
    
    # Case Details Table
    table_data = [
        [
            Paragraph("<b>Case Title:</b>", styles['Normal']), 
            Paragraph(case_data.get('title', 'N/A'), styles['Normal']),
            Paragraph("<b>Date of Report:</b>", styles['Normal']),
            Paragraph(datetime.now().strftime('%Y-%m-%d %H:%M'), styles['Normal'])
        ],
        [
            Paragraph("<b>Location:</b>", styles['Normal']),
            Paragraph(case_data.get('location', 'N/A'), styles['Normal']),
            Paragraph("<b>Incident Date:</b>", styles['Normal']),
            Paragraph(case_data.get('date', 'N/A'), styles['Normal'])
        ],
        [
            Paragraph("<b>Officer In Charge:</b>", styles['Normal']),
            Paragraph(case_data.get('officer', 'N/A'), styles['Normal']),
            Paragraph("<b>Status:</b>", styles['Normal']),
            Paragraph(case_data.get('status', 'Open').upper(), styles['Normal'])
        ]
    ]
    
    t = Table(table_data, colWidths=[100, 160, 100, 160])
    t.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), colors.HexColor('#f8fafc')),
        ('ALIGN', (0,0), (-1,-1), 'LEFT'),
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('LEFTPADDING', (0,0), (-1,-1), 8),
        ('RIGHTPADDING', (0,0), (-1,-1), 8),
        ('BOX', (0,0), (-1,-1), 1, colors.HexColor('#cbd5e1')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#e2e8f0')),
    ]))
    story.append(t)
    story.append(Spacer(1, 15))
    
    # Evidence & Witness Section
    story.append(Paragraph("CASE DESCRIPTION AND SUBMITTED EVIDENCE", styles['Heading2']))
    story.append(Spacer(1, 4))
    
    body_style = ParagraphStyle(
        'TableBody',
        parent=styles['Normal'],
        fontSize=9,
        leading=13
    )
    
    evidence_desc = case_data.get('description', '')
    evidence_items = case_data.get('evidence', '')
    witness_info = case_data.get('witness_details', '')
    
    meta_table_data = [
        [Paragraph("<b>Incident Narrative:</b>", body_style), Paragraph(evidence_desc, body_style)],
        [Paragraph("<b>Items of Evidence:</b>", body_style), Paragraph(evidence_items, body_style)],
        [Paragraph("<b>Witness Statements:</b>", body_style), Paragraph(witness_info, body_style)]
    ]
    meta_t = Table(meta_table_data, colWidths=[120, 400])
    meta_t.setStyle(TableStyle([
        ('VALIGN', (0,0), (-1,-1), 'TOP'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 6),
        ('TOPPADDING', (0,0), (-1,-1), 6),
        ('BOX', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
        ('INNERGRID', (0,0), (-1,-1), 0.5, colors.HexColor('#cbd5e1')),
    ]))
    story.append(meta_t)
    story.append(Spacer(1, 20))
    
    # AI Analysis & Drafts (Parsed Markdown)
    story.append(Paragraph("LEGAL FINDINGS AND DRAFTS (GENERATED)", styles['Heading2']))
    story.append(Spacer(1, 5))
    
    analysis_flowables = parse_markdown_to_reportlab(analysis_output, styles)
    story.extend(analysis_flowables)
    
    story.append(Spacer(1, 20))
    
    # Verification Signature Block
    sig_data = [
        [
            Paragraph("<b>Investigating Officer Signature:</b><br/><br/>___________________________", body_style),
            Paragraph("<b>Authorized Inspector Sign-Off:</b><br/><br/>___________________________", body_style)
        ]
    ]
    sig_t = Table(sig_data, colWidths=[260, 260])
    sig_t.setStyle(TableStyle([
        ('ALIGN', (0,0), (-1,-1), 'CENTER'),
        ('VALIGN', (0,0), (-1,-1), 'BOTTOM'),
        ('BOTTOMPADDING', (0,0), (-1,-1), 10),
    ]))
    
    story.append(sig_t)
    
    doc.build(story)
    pdf_bytes = buffer.getvalue()
    buffer.close()
    return pdf_bytes

def generate_docx_report(case_data: dict, analysis_output: str) -> bytes:
    doc = DocxDocument()
    
    # Layout Config
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        
    # Styles Setup
    style_normal = doc.styles['Normal']
    font = style_normal.font
    font.name = 'Arial'
    font.size = Pt(10.5)
    
    # Document Title
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.add_run("CRIMEGPT CASE REPORT & LEGAL ANALYSIS")
    title_run.font.name = 'Arial'
    title_run.font.size = Pt(18)
    title_run.bold = True
    title_run.font.color.rgb = colors.HexColor('#1e3a8a') # Blue
    
    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = subtitle.add_run("GENERATED VIA LAW ENFORCEMENT LEGAL PORTAL")
    sub_run.font.size = Pt(9)
    sub_run.italic = True
    sub_run.font.color.rgb = colors.HexColor('#64748b')
    
    doc.add_paragraph("-" * 60).alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Case Details Table
    table = doc.add_table(rows=3, cols=2)
    table.style = 'Table Grid'
    
    details = [
        ("Case Title", case_data.get('title', 'N/A')),
        ("Officer In Charge", case_data.get('officer', 'N/A')),
        ("Location", case_data.get('location', 'N/A')),
        ("Incident Date", case_data.get('date', 'N/A')),
        ("Current Status", case_data.get('status', 'Open').upper()),
        ("Report Date", datetime.now().strftime('%Y-%m-%d %H:%M'))
    ]
    
    cells = table.rows[0].cells
    cells[0].paragraphs[0].add_run(f"Case Title: {details[0][1]}").bold = True
    cells[1].paragraphs[0].add_run(f"Officer In Charge: {details[1][1]}")
    
    cells = table.rows[1].cells
    cells[0].paragraphs[0].add_run(f"Location: {details[2][1]}")
    cells[1].paragraphs[0].add_run(f"Incident Date: {details[3][1]}")
    
    cells = table.rows[2].cells
    cells[0].paragraphs[0].add_run(f"Current Status: {details[4][1]}").bold = True
    cells[1].paragraphs[0].add_run(f"Report Date: {details[5][1]}")
    
    doc.add_paragraph() # Spacer
    
    # Narratives Section
    h_narrative = doc.add_heading("Case Narratives & Evidence Details", level=2)
    h_narrative.runs[0].font.color.rgb = colors.HexColor('#0f766e')
    
    p_desc = doc.add_paragraph()
    p_desc.add_run("Incident Description: ").bold = True
    p_desc.add_run(case_data.get('description', ''))
    
    p_ev = doc.add_paragraph()
    p_ev.add_run("Evidence Gathered: ").bold = True
    p_ev.add_run(case_data.get('evidence', ''))
    
    p_wit = doc.add_paragraph()
    p_wit.add_run("Witness Statements: ").bold = True
    p_wit.add_run(case_data.get('witness_details', ''))
    
    doc.add_paragraph() # Spacer
    
    # AI Findings Sections
    h_analysis = doc.add_heading("AI Generated Analysis & Draft FIR", level=2)
    h_analysis.runs[0].font.color.rgb = colors.HexColor('#0f766e')
    
    lines = analysis_output.split('\n')
    for line in lines:
        line_strip = line.strip()
        if not line_strip:
            continue
            
        if line_strip.startswith('# '):
            p = doc.add_heading(line_strip[2:], level=1)
            p.runs[0].font.color.rgb = colors.HexColor('#1e3a8a')
        elif line_strip.startswith('## '):
            p = doc.add_heading(line_strip[3:], level=2)
            p.runs[0].font.color.rgb = colors.HexColor('#0f766e')
        elif line_strip.startswith('### '):
            p = doc.add_heading(line_strip[4:], level=3)
        elif line_strip.startswith('- ') or line_strip.startswith('* '):
            p = doc.add_paragraph(style='List Bullet')
            text = re.sub(r'\*\*(.*?)\*\*', r'\1', line_strip[2:])
            p.add_run(text)
        else:
            p = doc.add_paragraph()
            text = line_strip
            # Add simple formatting parse
            parts = re.split(r'(\*\*.*?\*\*)', text)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    p.add_run(part[2:-2]).bold = True
                else:
                    p.add_run(part)
                    
    doc.add_paragraph() # Spacer
    doc.add_paragraph()
    
    # Signature blocks
    p_sig = doc.add_paragraph()
    p_sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run_sig_left = p_sig.add_run("Investigating Officer: ___________________        ")
    run_sig_left.bold = True
    run_sig_right = p_sig.add_run("Counter-Signing Inspector: ___________________")
    run_sig_right.bold = True
    
    buffer = io.BytesIO()
    doc.save(buffer)
    docx_bytes = buffer.getvalue()
    buffer.close()
    return docx_bytes
